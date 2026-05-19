from __future__ import annotations

import argparse
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from mcp.server.fastmcp import FastMCP


DEFAULT_ALLOWED_ROOTS = [
    "/root/autodl-tmp/VTON/paper",
    "/root/autodl-tmp/VTON/log",
]


@dataclass(frozen=True)
class ServerConfig:
    allowed_roots: List[str]


def _normalize_abs(path: str) -> str:
    return str(Path(path).expanduser().resolve())


def _is_under(path: str, root: str) -> bool:
    p = Path(path)
    r = Path(root)
    try:
        p.relative_to(r)
        return True
    except Exception:
        return False


def _guard_path(cfg: ServerConfig, path: str) -> str:
    ap = _normalize_abs(path)
    if not os.path.isabs(ap):
        raise ValueError("path must be absolute")
    if not any(_is_under(ap, _normalize_abs(r)) for r in cfg.allowed_roots):
        raise ValueError(f"path not allowed: {ap}")
    return ap


def _open_docx(cfg: ServerConfig, path: str) -> Document:
    ap = _guard_path(cfg, path)
    if not ap.lower().endswith(".docx"):
        raise ValueError("only .docx is supported")
    if not os.path.exists(ap):
        raise FileNotFoundError(ap)
    return Document(ap)


def _save_docx(cfg: ServerConfig, doc: Document, path: str) -> str:
    ap = _guard_path(cfg, path)
    if not ap.lower().endswith(".docx"):
        raise ValueError("only .docx is supported")
    Path(ap).parent.mkdir(parents=True, exist_ok=True)
    doc.save(ap)
    return ap


def _paragraph_to_dict(p) -> Dict[str, Any]:
    return {
        "style": getattr(p.style, "name", ""),
        "text": p.text,
    }


def build_server(cfg: ServerConfig) -> FastMCP:
    mcp = FastMCP("VTON Word Tools", json_response=True)

    @mcp.tool()
    def docx_info(path: str) -> Dict[str, Any]:
        doc = _open_docx(cfg, path)
        styles = []
        seen = set()
        for s in doc.styles:
            n = getattr(s, "name", "")
            if not n or n in seen:
                continue
            seen.add(n)
            styles.append(n)
        return {
            "path": _normalize_abs(path),
            "paragraph_count": len(doc.paragraphs),
            "style_count": len(styles),
            "styles": styles[:200],
        }

    @mcp.tool()
    def docx_list_paragraphs(path: str, start: int = 0, limit: int = 50) -> Dict[str, Any]:
        doc = _open_docx(cfg, path)
        if start < 0:
            start = 0
        if limit < 1:
            limit = 1
        end = min(len(doc.paragraphs), start + limit)
        items = []
        for idx in range(start, end):
            items.append({"index": idx, **_paragraph_to_dict(doc.paragraphs[idx])})
        return {"path": _normalize_abs(path), "start": start, "end": end, "items": items}

    @mcp.tool()
    def docx_find(path: str, query: str, match_case: bool = False, limit: int = 200) -> Dict[str, Any]:
        doc = _open_docx(cfg, path)
        q = query if match_case else query.lower()
        hits = []
        for idx, p in enumerate(doc.paragraphs):
            text = p.text
            hay = text if match_case else text.lower()
            if q in hay:
                hits.append({"index": idx, "style": getattr(p.style, "name", ""), "text": text})
                if len(hits) >= limit:
                    break
        return {"path": _normalize_abs(path), "query": query, "match_case": match_case, "hits": hits}

    @mcp.tool()
    def docx_replace(path: str, old: str, new: str, match_case: bool = False) -> Dict[str, Any]:
        doc = _open_docx(cfg, path)
        replaced = 0
        for p in doc.paragraphs:
            if not p.runs:
                continue
            full = "".join(r.text for r in p.runs)
            if match_case:
                if old not in full:
                    continue
                updated = full.replace(old, new)
            else:
                lower_full = full.lower()
                lower_old = old.lower()
                if lower_old not in lower_full:
                    continue
                out = []
                i = 0
                while True:
                    j = lower_full.find(lower_old, i)
                    if j < 0:
                        out.append(full[i:])
                        break
                    out.append(full[i:j])
                    out.append(new)
                    i = j + len(old)
                updated = "".join(out)
            if updated != full:
                for r in p.runs:
                    r.text = ""
                p.runs[0].text = updated
                replaced += 1
        _save_docx(cfg, doc, path)
        return {"path": _normalize_abs(path), "replaced_paragraphs": replaced}

    @mcp.tool()
    def docx_set_paragraph_text(path: str, index: int, text: str) -> Dict[str, Any]:
        doc = _open_docx(cfg, path)
        if index < 0 or index >= len(doc.paragraphs):
            raise IndexError("paragraph index out of range")
        p = doc.paragraphs[index]
        if p.runs:
            for r in p.runs:
                r.text = ""
            p.runs[0].text = text
        else:
            p.add_run(text)
        _save_docx(cfg, doc, path)
        return {"path": _normalize_abs(path), "index": index, "style": getattr(p.style, "name", ""), "text": p.text}

    @mcp.tool()
    def docx_append_paragraph(path: str, text: str, style: Optional[str] = None) -> Dict[str, Any]:
        doc = _open_docx(cfg, path)
        p = doc.add_paragraph(text)
        if style:
            try:
                p.style = style
            except Exception:
                pass
        _save_docx(cfg, doc, path)
        return {"path": _normalize_abs(path), "index": len(doc.paragraphs) - 1, "style": getattr(p.style, "name", ""), "text": p.text}

    @mcp.tool()
    def docx_add_heading(path: str, text: str, level: int = 1) -> Dict[str, Any]:
        doc = _open_docx(cfg, path)
        lvl = max(0, min(int(level), 9))
        p = doc.add_heading(text, level=lvl)
        _save_docx(cfg, doc, path)
        return {"path": _normalize_abs(path), "index": len(doc.paragraphs) - 1, "style": getattr(p.style, "name", ""), "text": p.text}

    @mcp.tool()
    def docx_save_as(src_path: str, dst_path: str) -> Dict[str, Any]:
        src = _guard_path(cfg, src_path)
        dst = _guard_path(cfg, dst_path)
        if not src.lower().endswith(".docx") or not dst.lower().endswith(".docx"):
            raise ValueError("only .docx is supported")
        if not os.path.exists(src):
            raise FileNotFoundError(src)
        doc = Document(src)
        _save_docx(cfg, doc, dst)
        return {"src_path": src, "dst_path": dst}

    return mcp


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--transport", choices=["stdio", "streamable-http"], default="stdio")
    parser.add_argument("--host", default="127.0.0.1")
    parser.add_argument("--port", type=int, default=8000)
    parser.add_argument(
        "--allow-root",
        action="append",
        default=[],
        help="Add an allowed root directory (absolute path). Default allows /root/autodl-tmp/VTON/paper and /root/autodl-tmp/VTON/log.",
    )
    args = parser.parse_args()

    allowed = DEFAULT_ALLOWED_ROOTS[:]
    for r in args.allow_root:
        if r:
            allowed.append(r)
    cfg = ServerConfig(allowed_roots=allowed)
    mcp = build_server(cfg)

    if args.transport == "stdio":
        mcp.run(transport="stdio")
    else:
        mcp.run(transport="streamable-http", host=args.host, port=args.port)


if __name__ == "__main__":
    main()

